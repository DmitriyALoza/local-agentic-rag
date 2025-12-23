"""
Word (DOCX) Document Parser
Extracts text and headings from Word documents
"""

from pathlib import Path
from typing import List, Dict
from docx import Document


class WordParser:
    """Parser for Word (.docx) files"""

    def parse(self, file_path: str) -> List[Dict]:
        """
        Parse a Word document and extract structured content

        Args:
            file_path: Path to the .docx file

        Returns:
            List of dictionaries, each containing:
            - text: The extracted text content
            - section: Section identifier (e.g., heading or paragraph number)
            - metadata: Additional metadata about the content
        """
        doc = Document(file_path)
        chunks = []
        filename = Path(file_path).name

        current_section = "Document Start"
        section_counter = 0

        for para_num, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()

            if not text:
                continue

            # Check if this is a heading
            is_heading = paragraph.style.name.startswith('Heading')

            if is_heading:
                # Update current section to this heading
                current_section = text
                section_counter += 1

                chunks.append({
                    "text": text,
                    "section": f"Heading: {text}",
                    "metadata": {
                        "filename": filename,
                        "paragraph_number": para_num,
                        "style": paragraph.style.name,
                        "content_type": "heading",
                        "section_number": section_counter
                    }
                })
            else:
                # Regular paragraph
                chunks.append({
                    "text": text,
                    "section": current_section,
                    "metadata": {
                        "filename": filename,
                        "paragraph_number": para_num,
                        "style": paragraph.style.name,
                        "content_type": "paragraph",
                        "parent_section": current_section
                    }
                })

        # Extract tables
        for table_num, table in enumerate(doc.tables, start=1):
            table_text_parts = []

            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text_parts.append(row_text)

            if table_text_parts:
                table_text = "\n".join(table_text_parts)
                chunks.append({
                    "text": table_text,
                    "section": f"Table {table_num}",
                    "metadata": {
                        "filename": filename,
                        "content_type": "table",
                        "table_number": table_num,
                        "rows": len(table.rows),
                        "columns": len(table.columns)
                    }
                })

        return chunks

    def get_metadata(self, file_path: str) -> Dict:
        """
        Get metadata about the Word document

        Args:
            file_path: Path to the .docx file

        Returns:
            Dictionary containing file metadata
        """
        doc = Document(file_path)
        path = Path(file_path)

        # Count headings
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]

        return {
            "filename": path.name,
            "file_type": "docx",
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "total_headings": len(headings),
            "file_size_bytes": path.stat().st_size
        }
